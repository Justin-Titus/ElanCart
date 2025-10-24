# This script generates ElanCart_Documentation.docx using python-docx.
# Run: pip install python-docx
"""
generate_elancart_doc.py

Creates a professional internship project report for the ElanCart React front-end.
Only standard library + python-docx are required.

Usage: python generate_elancart_doc.py

The script exposes generate_document(output_path, metadata, images=None)
"""

import os
import sys
import textwrap
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    print("The 'python-docx' package is required. Install it with: pip install python-docx")
    raise


# -------------------------
# Utility helpers
# -------------------------

def _sanitize(s):
    """Remove problematic control characters for XML/Word."""
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    # Remove C0 control chars except CR, LF and TAB
    import re

    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", s)


def _create_page_number_field():
    """Return an OxmlElement that represents a PAGE field for page numbers."""
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    return fld


# -------------------------
# Document construction helpers
# -------------------------

def create_document():
    doc = Document()

    # Set 1" margins for clean layout
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Set default Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    return doc


def add_cover_page(doc, metadata):
    # Cover: large centered title, subtitle, company, author/trainer/dates and logo placeholder
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(_sanitize(metadata.get('project_title', 'ElanCart')))
    run.bold = True
    run.font.size = Pt(26)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    srun = subtitle.add_run('React Internship Final Project')
    srun.italic = True
    srun.font.size = Pt(14)

    doc.add_paragraph()

    company_para = doc.add_paragraph()
    company_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    c_run = company_para.add_run(_sanitize(metadata.get('company_name', 'Softek Square Private Limited')))
    c_run.font.size = Pt(12)

    doc.add_paragraph()

    prepared = doc.add_paragraph()
    prepared.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_run = prepared.add_run('Prepared by\n')
    p_run.bold = True
    p_run.font.size = Pt(11)

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info.add_run(_sanitize(metadata.get('intern_name', 'Intern Name Placeholder')) + '\n')
    info.add_run('Trainer / Mentor: ' + _sanitize(metadata.get('trainer_name', 'Trainer Placeholder')) + '\n')
    info.add_run('Internship Dates: ' + _sanitize(metadata.get('start_date', 'Start')) + ' — ' + _sanitize(metadata.get('end_date', 'End')))

    doc.add_paragraph()
    logo = doc.add_paragraph()
    logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    logo_run = logo.add_run('[Company Logo Placeholder]')
    logo_run.italic = True
    logo_run.font.size = Pt(10)

    # small generation note
    gen = doc.add_paragraph()
    gen.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    gen.add_run('Generated on: ' + datetime.now().strftime('%B %d, %Y'))

    doc.add_page_break()


def add_document_metadata(doc, metadata):
    doc.add_heading('Document Metadata', level=1)
    meta = {
        'Project': metadata.get('project_title', 'ElanCart'),
        'Version': metadata.get('version', '1.0'),
        'Created': datetime.now().strftime('%Y-%m-%d %H:%M'),
        'Author': metadata.get('intern_name', 'Intern Name Placeholder'),
    }
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Key'
    hdr[1].text = 'Value'
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for k, v in meta.items():
        row = table.add_row().cells
        row[0].text = _sanitize(k)
        row[1].text = _sanitize(v)
    doc.add_paragraph()


def add_toc_placeholder(doc):
    doc.add_heading('Table of Contents', level=1)
    entries = [
        '1. Project Overview',
        '2. Objective / Problem Statement',
        '3. Technologies & Libraries',
        '4. Project Structure & Key Files',
        '5. How to Run Locally',
        '6. Live Website',
        '7. Features & User Flows',
        '8. Extending & Next Steps',
        '9. Appendix',
        '10. Sign-off / Verification',
    ]
    for e in entries:
        p = doc.add_paragraph()
        p.add_run(e)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def add_section_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(_sanitize(text))
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    p.paragraph_format.space_after = Pt(6)


def add_bulleted_list(doc, items):
    for it in items:
        p = doc.add_paragraph(_sanitize(it), style='List Bullet')
        p.paragraph_format.space_after = Pt(2)


def add_numbered_list(doc, items):
    for it in items:
        p = doc.add_paragraph(_sanitize(it), style='List Number')
        p.paragraph_format.space_after = Pt(2)


def add_kv_table(doc, data, col_names=('Path', 'Description')):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = col_names[0]
    hdr[1].text = col_names[1]
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    # Add rows
    for key, val in data:
        r = table.add_row().cells
        r[0].text = _sanitize(key)
        r[1].text = _sanitize(val)
    doc.add_paragraph()


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    run = p.add_run(_sanitize(code_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_project_structure_table(doc):
    doc.add_heading('Project Structure & Key Files', level=1)
    files = [
        ('src/main.jsx', 'Application entry and root layout; wires up contexts and routing.'),
        ('src/components/layout', 'Header, Footer, Layout components for scaffolding.'),
        ('src/components/products', 'Product grid, cards, filters, banners for discovery.'),
        ('src/contexts', 'CartContext.jsx, ProductContext.jsx, etc. for global state.'),
    ]
    add_kv_table(doc, files, col_names=('Path', 'Role / Description'))


def add_how_to_run(doc):
    doc.add_heading('How to Run Locally', level=1)
    add_paragraph(doc, 'Prerequisites: Node.js 16+ and npm')
    steps = [
        '1. From project root: npm install',
        '2. Start development server: npm run dev',
        '3. Build for production: npm run build',
        '4. Preview production build: npm run preview',
    ]
    add_numbered_list(doc, steps)
    add_paragraph(doc, 'Note: Update API base URL in src/axios.js to point to your backend before running production builds.', italic=True)


def add_live_website(doc):
    doc.add_heading('Live Website', level=1)
    add_paragraph(doc, 'Production demo (ShopEasy): https://elancart.netlify.app/')
    add_paragraph(doc, 'Deployment: The site is suitable for static hosting (Netlify/Vercel). The deploy pipeline typically builds with Vite and publishes the dist output to the hosting provider.')


def add_features_and_flows(doc):
    doc.add_heading('Features & User Flows', level=1)

    # Discovery
    doc.add_heading('Discovery / Browsing', level=2)
    add_bulleted_list(doc, [
        'Responsive product grid with quick filters and categories.',
        'Search and sorting capabilities for efficient discovery.',
        'Product cards with CTA: add-to-cart, view details, and favourite.'
    ])

    # PDP
    doc.add_heading('Product Detail Page (PDP)', level=2)
    add_bulleted_list(doc, [
        'Image gallery and detailed product information.',
        'Attributes, pricing, and stock status presented clearly.',
        'Add to cart and buy-now actions; favourites and sharing options.'
    ])

    # Cart
    doc.add_heading('Cart & Checkout', level=2)
    add_bulleted_list(doc, [
        'Cart preview with quantity management and subtotal calculations.',
        'Checkout flow includes address selection and order review.',
        'Payment gateway placeholders show where integration should be added.'
    ])

    # Favourites/profile
    doc.add_heading('Favourites & Profile', level=2)
    add_bulleted_list(doc, [
        'Save items for later in a favourites list.',
        'Profile scaffolding for future authentication and order history.'
    ])

    # Accessibility/responsiveness
    doc.add_heading('Accessibility & Responsiveness', level=2)
    add_paragraph(doc, 'Accessibility: semantic HTML, keyboard navigability, and ARIA attributes where applicable. Responsive: mobile-first styles and breakpoints for common devices.')


def add_extending_next_steps(doc):
    doc.add_heading('Extending & Next Steps', level=1)
    items = [
        'Backend & API',
        'Persistence',
        'Authentication',
        'Payment integration',
        'Tests',
        'CI/CD',
        'Performance',
    ]
    # Each with 1-2 practical suggestions
    details = {
        'Backend & API': ['Design REST endpoints for products, carts, orders.', 'Implement server-side validation and rate limiting.'],
        'Persistence': ['Persist cart and favourites server-side and associate with user accounts.', 'Add migration and backup strategies.'],
        'Authentication': ['Add JWT-based auth or OAuth flows.', 'Protect sensitive routes and add role checks.'],
        'Payment integration': ['Integrate Stripe or PayPal via secure server-side endpoints.', 'Avoid handling raw card data on the client.'],
        'Tests': ['Add unit tests for key components (ProductCard, CartItem).', 'Use Cypress for end-to-end checkout flow tests.'],
        'CI/CD': ['Add GitHub Actions to run lint and tests on PRs.', 'Deploy successful builds to Netlify with preview links.'],
        'Performance': ['Lazy-load images and use code-splitting for large routes.', 'Measure with Lighthouse and address high-impact findings.'],
    }
    for it in items:
        doc.add_heading(it, level=2)
        add_bulleted_list(doc, details.get(it, []))

    # Low-risk developer tasks
    doc.add_heading('Low-risk developer tasks', level=2)
    add_bulleted_list(doc, [
        'Add README sections for local development and environment variables.',
        'Create .env.example to document required environment variables.',
        'Add unit tests for critical components and a small CI workflow.'
    ])


def add_appendix(doc, images=None):
    doc.add_heading('Appendix: Screenshots / Assets', level=1)
    if not images:
        add_paragraph(doc, 'No local images were provided. To include screenshots or assets, pass an images dict to the script.')
        # leave a blank line for spacing
        add_paragraph(doc, '')
    else:
        for caption, path in images.items():
            add_paragraph(doc, caption, bold=True)
            if os.path.exists(path):
                try:
                    # Resize to max width 6 inches while keeping aspect ratio
                    doc.add_picture(path, width=Inches(6))
                except Exception as e:
                    add_paragraph(doc, f'Unable to insert image: {e}')
            else:
                add_paragraph(doc, f'Image not found: {path}', italic=True)
    doc.add_paragraph()

    # Add .env.example code block
    add_paragraph(doc, '.env.example', bold=True)
    env_example = 'REACT_APP_API_URL=https://api.example.com\nREACT_APP_OTHER_KEY=value'
    add_code_block(doc, env_example)

    # Sample API endpoints
    doc.add_heading('Sample API Endpoints', level=2)
    add_bulleted_list(doc, [
        'GET /products — List products with filters',
        'GET /products/:id — Product detail',
        'POST /auth/login — Authenticate user',
    ])


def add_signoff(doc, metadata):
    doc.add_page_break()
    doc.add_heading('Sign-off / Verification', level=1)
    add_paragraph(doc, 'Trainer / Mentor: ' + _sanitize(metadata.get('trainer_name', 'Trainer Placeholder')))
    add_paragraph(doc, 'Company: ' + _sanitize(metadata.get('company_name', 'Softek Square Private Limited')))
    add_paragraph(doc, 'Submission Date: ' + datetime.now().strftime('%B %d, %Y'))
    add_paragraph(doc, 'Signature: ____________________________')


def add_header_and_footer(doc, metadata):
    # Header: left project name, right page number
    section = doc.sections[0]
    header = section.header
    # Clear default
    if header.is_linked_to_previous:
        header.is_linked_to_previous = False
    header_p = header.paragraphs[0]
    header_p.text = ''
    left_run = header_p.add_run(_sanitize(metadata.get('project_title', 'ElanCart')))
    left_run.bold = True
    header_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Right-aligned page number
    page_run = header_p.add_run()
    page_run._r.append(_create_page_number_field())
    header_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Footer: Confidential note + page numbers
    footer = section.footer
    if footer.is_linked_to_previous:
        footer.is_linked_to_previous = False
    f_p = footer.paragraphs[0]
    f_p.text = 'Confidential — Internship Project'
    # Add a space and page number
    run = f_p.add_run('  ')
    run._r.append(_create_page_number_field())
    f_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# -------------------------
# Main generation function
# -------------------------

def generate_document(output_path, metadata, images=None):
    """Generate the ElanCart project documentation and save to output_path.

    metadata keys expected: project_title, intern_name, trainer_name, company_name, start_date, end_date, version
    images: optional dict mapping caption -> local file path
    """
    try:
        doc = create_document()

        # Cover
        add_cover_page(doc, metadata)

        # Metadata
        add_document_metadata(doc, metadata)

        # TOC placeholder
        add_toc_placeholder(doc)

        # Project Overview (use exact paragraph provided)
        doc.add_heading('Project Overview', level=1)
        project_overview = (
            'ElanCart is a responsive e-commerce front-end built with React that showcases a small-to-medium online store experience. '
            'It provides product browsing, filtering, product detail pages, a shopping cart, favourites, and checkout flows (UI only). '
            'The project demonstrates component-driven UI, context-based state management, modular CSS, and integration points for REST APIs and payment gateways.'
        )
        add_paragraph(doc, project_overview)

        # Objective / Problem Statement
        doc.add_heading('Objective / Problem Statement', level=1)
        # Provided bullets (verbatim from earlier content)
        objectives = [
            'A clear product discovery experience with filters, quick categories, and banners.',
            'Detailed product pages with images, attributes, and add-to-cart functionality.',
            'A simple, accessible cart and checkout UI (payment gateway integration points included as sample pages).',
            'Reusable components and context providers to keep UI state consistent (cart, user, locale, favourites).'
        ]
        add_bulleted_list(doc, objectives)
        # Expand into 2-3 formal sentences
        obj_expansion = (
            'The objectives of this internship project were to design and build a maintainable, responsive front-end that models common e-commerce flows. '
            'The implementation focuses on usability, component reusability and a clear separation between presentation and data services, making it suitable for integration with a future backend. '
            'Emphasis was placed on accessibility and cross-device compatibility to ensure a consistent user experience.'
        )
        add_paragraph(doc, obj_expansion)

        # Technologies & Libraries
        doc.add_heading('Technologies & Libraries', level=1)
        techs = [
            ('React', 'Core UI library used to build components and pages.'),
            ('Context API', 'Global state management for cart, user and preferences.'),
            ('CSS Modules', 'Scoped styling for components to avoid global collisions.'),
            ('Axios', 'HTTP client used by service adapters to call APIs.'),
            ('Vite / npm', 'Build and development tooling used to run and build the app.'),
            ('Payment placeholders', 'Pages and UI hooks indicating where payment integrations should be connected.')
        ]
        for name, desc in techs:
            add_paragraph(doc, f'{name} — {desc}')

        # Project Structure
        add_project_structure_table(doc)

        # How to run
        add_how_to_run(doc)

        # Live website
        add_live_website(doc)

        # Features & flows
        add_features_and_flows(doc)

        # Extending & Next Steps
        add_extending_next_steps(doc)

        # Appendix (images and extras)
        add_appendix(doc, images)

        # Sign-off
        add_signoff(doc, metadata)

        # Header & footer
        add_header_and_footer(doc, metadata)

        # Save
        try:
            doc.save(output_path)
            print(f'Saved: {output_path}')
        except Exception as e:
            print('Error saving document:', e)
            raise

    except Exception as exc:
        print('Failed to generate document:', exc)
        raise


# -------------------------
# CLI demo
# -------------------------
if __name__ == '__main__':
    metadata = {
        'project_title': 'ElanCart',
        'intern_name': 'Justin Titus J.',
        'trainer_name': 'Trainer Placeholder',
        'company_name': 'Softek Square Private Limited',
        'start_date': 'June 2025',
        'end_date': 'Oct 2025',
        'version': '1.0',
    }

    output = 'ElanCart_Documentation.docx'

    # Optionally pass images as a dict: {'Homepage screenshot': 'assets/home.png'}
    images = None

    generate_document(output, metadata, images=images)

